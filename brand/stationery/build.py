#!/usr/bin/env python3
"""ṚTAM stationery driver — renders the print masters, previews, and the Word
typing template from the HTML sources in this directory.

    python3 brand/stationery/build.py            # everything
    python3 brand/stationery/build.py letterhead # one piece

Print masters are vector PDFs via the shared Chromium renderer (page='css' so
each document's @page rule controls the sheet). The DOCX typing template
carries the head/footer as images in the Word header/footer layer (Oxford
pattern — typists cannot move them); its body face falls back gracefully on
machines without the brand fonts.
"""
from __future__ import annotations

import sys
from pathlib import Path

HERE = Path(__file__).resolve().parent
BRAND = HERE.parent
TOOLS = BRAND / "tools"
DIST = HERE / "dist"
PREVIEWS = BRAND / "exports" / "stationery"
sys.path.insert(0, str(TOOLS))

from render_html_to_png import render as render_html  # noqa: E402

MM = 72 / 25.4

# (source html, master pdf, preview png) — page size comes from each @page rule
PIECES = {
    "letterhead":      ("letterhead.html",      "letterhead.pdf",      "letterhead-preview.png"),
    "letterhead-mono": ("letterhead-mono.html", "letterhead-mono.pdf", "letterhead-mono-preview.png"),
    "seal-round":      ("seal-chakra-round.html", "seal-chakra-round.pdf", None),
    "address-stamp":   ("address-stamp.html",     "address-stamp.pdf",     None),
    "receipt":         ("receipt-a5.html",      "receipt-a5.pdf",      "receipt-a5-preview.png"),
    "receipt-cover":   ("receipt-cover-a5.html","receipt-cover-a5.pdf","receipt-cover-a5-preview.png"),
    "labels":          ("labels-st8.html",      "labels-st8.pdf",      "labels-st8-preview.png"),
    "labels-test":     ("labels-alignment-test.html", "labels-alignment-test.pdf", None),
}


def build_piece(name: str) -> None:
    src, pdf, png = PIECES[name]
    html = HERE / src
    if not html.exists():
        print(f"  skip  {name} ({src} not present yet)")
        return
    render_html(html, DIST / pdf, 840, 2, "css")
    print(f"  pdf   dist/{pdf}")
    if png:
        render_html(html, PREVIEWS / png, 840, 2)
        print(f"  png   exports/stationery/{png}")


def build_us_editions() -> None:
    """US Letter editions of the letterhead (215.9×279.4mm) for printing in
    the US — the founder's HP tray holds Letter, and A4 masters printed on
    Letter at 100% lose the bottom 17.6mm (the fourth footer line). Generated
    from the A4 sources by page-size substitution; the footer is
    bottom-anchored so the design re-flows correctly."""
    for src_name, out_stem in [("letterhead.html", "letterhead-us"),
                               ("letterhead-mono.html", "letterhead-us-mono")]:
        src = (HERE / src_name).read_text()
        us = (src
              .replace("@page { size: A4; margin: 0; }",
                       "@page { size: 215.9mm 279.4mm; margin: 0; }")
              .replace("</head>",
                       "<style>.sheet { width: 215.9mm; height: 278.9mm; }</style></head>"))
        tmp = HERE / f"{out_stem}.html"
        tmp.write_text(us)
        render_html(tmp, DIST / f"{out_stem}.pdf", 860, 2, "css")
        tmp.unlink()
        print(f"  pdf   dist/{out_stem}.pdf")


def build_receipt_color() -> None:
    """The colored donor-facing edition of the receipt: generated from
    receipt-a5.html by swapping the two single-ink cuts for the gold-hub cuts
    (one source, two inks — the letterhead press/mono pattern, mechanized so
    the editions can never drift apart)."""
    src = (HERE / "receipt-a5.html").read_text()
    color = (src
             .replace("rtam-chakra-mono.svg", "rtam-chakra-day.svg")
             .replace("rtambhareshvara-mandir-lockup-devanagari-led-charcoal.svg",
                      "rtambhareshvara-mandir-lockup-devanagari-led.svg")
             .replace("दान-रसीद — A5 landscape leaf (NCR duplicate book)",
                      "दान-रसीद — A5 landscape leaf (colour edition)"))
    tmp = HERE / "receipt-a5-color.html"
    tmp.write_text(color)
    render_html(tmp, DIST / "receipt-a5-color.pdf", 840, 2, "css")
    print("  pdf   dist/receipt-a5-color.pdf")
    render_html(tmp, PREVIEWS / "receipt-a5-color-preview.png", 840, 2)
    print("  png   exports/stationery/receipt-a5-color-preview.png")


def build_docx() -> None:
    """letterhead.docx — A4, margins matched to the sheet (25/20/58/28mm), the
    head and footer as strip images from a 576-dpi render of the master."""
    try:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm, Pt
    except ImportError:
        print("  skip  letterhead.docx (python-docx not installed — "
              "pip install --user python-docx, or use the PDF masters)")
        return
    from PIL import Image

    strip_src = DIST / "_letterhead-600.png"
    render_html(HERE / "letterhead.html", strip_src, 840, 6)  # ≈576 dpi
    img = Image.open(strip_src)
    pw, ph = img.size                       # rendered sheet incl. screen chrome
    # The @media screen desk surround pads the sheet: locate the white sheet by
    # scanning for the page box (the sheet is the centered 210mm-wide region).
    # Screen render: body pad 40px top at scale 6 = 240px; sheet width 210mm
    # at 96dpi*6 = 4762px.
    sheet_w = round(210 / 25.4 * 96 * 6)
    x0 = (pw - sheet_w) // 2
    y0 = 40 * 6
    px_per_mm = sheet_w / 210.0

    def crop(x_mm, y_mm, w_mm, h_mm, out: Path):
        box = (round(x0 + x_mm * px_per_mm), round(y0 + y_mm * px_per_mm),
               round(x0 + (x_mm + w_mm) * px_per_mm), round(y0 + (y_mm + h_mm) * px_per_mm))
        img.crop(box).save(out)

    head_png = DIST / "_docx-head.png"
    foot_png = DIST / "_docx-foot.png"
    crop(25, 18, 165, 37, head_png)         # chakra + lockup + refline
    crop(25, 264, 165, 26, foot_png)        # footer block incl. its rule

    doc = docx.Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.left_margin, sec.right_margin = Mm(25), Mm(20)
    sec.top_margin, sec.bottom_margin = Mm(58), Mm(28)
    sec.header_distance, sec.footer_distance = Mm(18), Mm(7)

    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hp.add_run().add_picture(str(head_png), width=Mm(165))

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_after = Pt(0)
    fp.add_run().add_picture(str(foot_png), width=Mm(165))

    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"            # graceful fallback body face
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.4
    doc.add_paragraph("")                   # the letter begins here

    out = DIST / "letterhead.docx"
    doc.save(out)
    for tmp in (strip_src, head_png, foot_png):
        tmp.unlink(missing_ok=True)
    print("  docx  dist/letterhead.docx")


def main(argv: list[str]) -> int:
    only = argv[1] if len(argv) > 1 else None
    DIST.mkdir(exist_ok=True)
    PREVIEWS.mkdir(parents=True, exist_ok=True)
    for name in PIECES:
        if only and only != name:
            continue
        build_piece(name)
    if only in (None, "receipt", "receipt-color"):
        build_receipt_color()
    if only in (None, "letterhead", "letterhead-us"):
        build_us_editions()
    if only in (None, "docx", "letterhead"):
        build_docx()
    print("stationery build: done")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
